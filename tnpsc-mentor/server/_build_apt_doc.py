#!/usr/bin/env python3
"""Build a Word doc of Aptitude PYQs, organised by YEAR then GROUP.
Group 1 = category 'pyq', Group 2 = category 'pyq2'.
Powers (^n) -> real superscripts. Fractions (a/b) -> stacked OMML fractions.
Markdown pipe-tables -> Word tables. Figure images embedded (Group 2 crops
from the bucket; Group 1 reconstructed figures where derivable).
Usage: python _build_apt_doc.py <input.json> <output.docx> "<title>"
"""
import json, re, sys, os
from xml.sax.saxutils import escape
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

inp, outp = sys.argv[1], sys.argv[2]
title = sys.argv[3] if len(sys.argv) > 3 else "TNPSC Aptitude — Previous Year Questions"

rows = json.load(open(inp, encoding="utf-8"))
GROUP_LABEL = {"pyq": "Group 1", "pyq2": "Group 2"}

# url -> local cached file
MANIFEST = {}
if os.path.exists("_imgcache/manifest.json"):
    MANIFEST = json.load(open("_imgcache/manifest.json", encoding="utf-8"))

# Single left-to-right scanner. POWER before FRACTION so an exponent fraction
# (e.g. x^{3/2}) is consumed by the power and not stacked.
NUMTOK = r"(?:\([^()]*\)|\d+(?:\.\d+)?)"
TOKEN = re.compile(
    r"\^(?P<exp>\([^)]*\)|\{[^}]*\}|[+\-]?[0-9A-Za-z]+)"
    rf"|(?P<num>{NUMTOK})\s*/\s*(?P<den>{NUMTOK})"
)
TABLE_LINE = re.compile(r"^\s*\|.*\|\s*$")


def strip_parens(s):
    s = s.strip()
    if s.startswith("(") and s.endswith(")"):
        inner, depth = s[1:-1], 0
        for ch in inner:
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth -= 1
                if depth < 0:
                    return s
        return inner.strip()
    return s


def omml_fraction(num, den):
    num, den = escape(strip_parens(num)), escape(strip_parens(den))
    xml = (
        f'<m:oMath {nsdecls("m", "w")}>'
        f'<m:f><m:num><m:r><m:t xml:space="preserve">{num}</m:t></m:r></m:num>'
        f'<m:den><m:r><m:t xml:space="preserve">{den}</m:t></m:r></m:den></m:f>'
        f'</m:oMath>'
    )
    return parse_xml(xml)


def add_text_runs(p, text, base_size=11, bold=False):
    """fractions -> OMML, powers -> superscript, newlines -> line breaks."""
    if text is None:
        text = ""

    def plain(s):
        if s:
            r = p.add_run(s); r.font.size = Pt(base_size); r.bold = bold

    for li, line in enumerate(str(text).split("\n")):
        if li > 0:
            p.add_run().add_break()
        pos = 0
        for m in TOKEN.finditer(line):
            plain(line[pos:m.start()])
            if m.group("exp") is not None:
                exp = m.group("exp")
                if exp[:1] in "({" and exp[-1:] in ")}":
                    exp = exp[1:-1]
                r = p.add_run(exp); r.font.size = Pt(base_size); r.bold = bold
                r.font.superscript = True
            else:
                p._p.append(omml_fraction(m.group("num"), m.group("den")))
            pos = m.end()
        plain(line[pos:])


def split_blocks(text):
    """Split question text into ('text', str) and ('table', [[cells]]) blocks."""
    blocks, buf, tbl = [], [], []

    def flush_text():
        if any(l.strip() for l in buf):
            blocks.append(("text", "\n".join(buf).strip("\n")))
        buf.clear()

    def flush_tbl():
        if tbl:
            blocks.append(("table", list(tbl)))
            tbl.clear()

    for line in str(text or "").split("\n"):
        if TABLE_LINE.match(line):
            flush_text()
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            tbl.append(cells)
        else:
            flush_tbl()
            buf.append(line)
    flush_text(); flush_tbl()
    return blocks


def add_table(doc, tbl):
    ncol = max(len(r) for r in tbl)
    t = doc.add_table(rows=len(tbl), cols=ncol)
    t.style = "Table Grid"
    t.autofit = True
    for i, row in enumerate(tbl):
        for j in range(ncol):
            cell = t.cell(i, j)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_text_runs(para, row[j] if j < len(row) else "", base_size=10)


def embed(doc, local_path, width_in=2.2):
    if local_path and os.path.exists(local_path):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        try:
            p.add_run().add_picture(local_path, width=Inches(width_in))
            return True
        except Exception as e:
            sys.stderr.write(f"img fail {local_path}: {e}\n")
    return False


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Nirmala UI"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = rpr.makeelement(qn("w:rFonts"), {})
    rpr.append(rfonts)
for attr in ("w:ascii", "w:hAnsi", "w:cs"):
    rfonts.set(qn(attr), "Nirmala UI")

doc.add_heading(title, level=0)

years = OrderedDict()
for r in rows:
    years.setdefault(r.get("year"), {}).setdefault(r["category"], []).append(r)

OPT_KEYS = [("A", "option_a"), ("B", "option_b"), ("C", "option_c"),
            ("D", "option_d"), ("E", "option_e")]

total = embedded = 0
for y in sorted(years, key=lambda v: (v is None, v)):
    doc.add_heading(f"Year {y}" if y else "Year — (untagged)", level=1)
    for cat in ("pyq", "pyq2"):
        grows = years[y].get(cat)
        if not grows:
            continue
        doc.add_heading(GROUP_LABEL[cat], level=2)
        n = 0
        for r in grows:
            n += 1
            total += 1
            has_img = bool(r.get("images"))
            # --- question text (with tables) ---
            qtext = r.get("question_text") or ""
            if has_img:
                # a real figure image replaces the inline grid / "[Figure: ...]" note
                qtext = "\n".join(
                    ln for ln in qtext.split("\n")
                    if not re.match(r"^\s*\[figure", ln, re.I)
                )
            blocks = split_blocks(qtext) or [("text", "")]
            first = True
            for kind, content in blocks:
                if kind == "table" and has_img:
                    continue  # image already shows the grid
                if kind == "text":
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8) if first else Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    if first:
                        p.add_run(f"{n}. ").bold = True
                    add_text_runs(p, content)
                else:
                    if first:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(8)
                        p.add_run(f"{n}. ").bold = True
                    add_table(doc, content)
                first = False
            if first:  # nothing rendered (all blocks suppressed) — still show number
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.add_run(f"{n}. ").bold = True

            # --- Tamil ---
            ta = r.get("question_text_ta")
            if ta and ta.strip():
                pt = doc.add_paragraph()
                pt.paragraph_format.space_after = Pt(2)
                add_text_runs(pt, ta)

            # --- figure images ---
            for url in (r.get("images") or []):
                if embed(doc, MANIFEST.get(url), 2.4):
                    embedded += 1

            # --- options ---
            optimg = r.get("option_images") or {}
            for letter, key in OPT_KEYS:
                val = r.get(key)
                has_img = letter in optimg
                if (val is None or str(val).strip() == "") and not has_img:
                    continue
                po = doc.add_paragraph()
                po.paragraph_format.left_indent = Pt(18)
                po.paragraph_format.space_after = Pt(0)
                po.add_run(f"({letter}) ").bold = True
                if val and str(val).strip():
                    add_text_runs(po, val)
                if has_img:
                    local = MANIFEST.get(optimg[letter])
                    if local and os.path.exists(local):
                        try:
                            po.add_run(" ").add_picture(local, height=Inches(0.5))
                            embedded += 1
                        except Exception as e:
                            sys.stderr.write(f"optimg fail: {e}\n")

            # --- answer + tags ---
            pa = doc.add_paragraph()
            pa.paragraph_format.space_after = Pt(4)
            pa.add_run("Answer: ").bold = True
            av = pa.add_run(str(r.get("correct_answer") or "-"))
            av.bold = True
            av.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            tagbits = [GROUP_LABEL[cat]]
            if r.get("year"):
                tagbits.append(str(r["year"]))
            if r.get("aptitude_type"):
                tagbits.append(str(r["aptitude_type"]).capitalize())
            if r.get("topic"):
                tagbits.append(str(r["topic"]))
            tr = pa.add_run("    [" + " · ".join(tagbits) + "]")
            tr.italic = True; tr.font.size = Pt(9)
            tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(outp)
print(f"Saved {outp} — {total} questions, {embedded} images embedded")
