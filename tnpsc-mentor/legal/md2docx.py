"""Minimal Markdown -> .docx converter for the TNPSC Mentor legal docs.

Handles: # / ## / ### headings, blockquotes (>), bullet lists (-),
horizontal rules (---), and inline **bold**, `code`, and [text](url) links.
Good enough for these structured policy documents; not a general MD engine.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))")
LINK = re.compile(r"\[(.+?)\]\((.+?)\)")


def add_runs(paragraph, text):
    """Split a line into runs, applying bold / monospace / link-text styling."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
        elif part.startswith("["):
            m = LINK.match(part)
            run = paragraph.add_run(m.group(1) if m else part)
        else:
            paragraph.add_run(part)


def convert(md_path, docx_path):
    doc = Document()
    # Base body font.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = open(md_path, encoding="utf-8").read().splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.strip() == "---":
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)  # title style
        elif line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line.lstrip(">").strip())
        elif line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line.lstrip()[2:])
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    for md, dx in [(sys.argv[1], sys.argv[2]), (sys.argv[3], sys.argv[4])]:
        convert(md, dx)
